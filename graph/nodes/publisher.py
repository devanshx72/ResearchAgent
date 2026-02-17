"""
Publisher/Exporter Node - Exports final article to specified format.
"""

import os
from datetime import datetime
from typing import Dict
from graph.state import ResearchState


def publisher_node(state: ResearchState) -> Dict:
    """
    Export the final article to the specified format.
    
    Supports: markdown (.md), docx (.docx), notion (optional)
    
    Args:
        state: Current research state
        
    Returns:
        Updated state with export_path
    """
    print("[Publisher] Publishing article...")
    
    final_article = state.get("final_article", "")
    sources = state.get("sources", [])
    export_format = state.get("export_format", "markdown")
    query = state.get("query", "research")
    
    # Create outputs directory if it doesn't exist
    output_dir = "outputs"
    os.makedirs(output_dir, exist_ok=True)
    
    # Generate filename with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_query = "".join(c if c.isalnum() or c in (' ', '_') else '_' for c in query)
    safe_query = safe_query.replace(' ', '_')[:50]  # Limit length
    
    export_path = ""
    
    if export_format == "markdown" or export_format == "md":
        export_path = _export_markdown(final_article, sources, safe_query, timestamp, output_dir)
    elif export_format == "docx":
        export_path = _export_docx(final_article, sources, safe_query, timestamp, output_dir)
    elif export_format == "notion":
        export_path = _export_notion(final_article, sources, safe_query)
    else:
        print(f"Warning: Unknown format '{export_format}', defaulting to markdown")
        export_path = _export_markdown(final_article, sources, safe_query, timestamp, output_dir)
    
    print(f"[Publisher] Article published: {export_path}")
    
    return {
        "export_path": export_path
    }


def _export_markdown(article: str, sources: list, query: str, timestamp: str, output_dir: str) -> str:
    """Export article as markdown file."""
    filename = f"{query}_{timestamp}.md"
    filepath = os.path.join(output_dir, filename)
    
    # Add sources section if not already present
    if "## Sources" not in article and "## References" not in article:
        article += "\n\n---\n\n## Sources\n\n"
        article += "\n".join([f"{i+1}. {url}" for i, url in enumerate(sources)])
    
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(article)
    
    return filepath


def _export_docx(article: str, sources: list, query: str, timestamp: str, output_dir: str) -> str:
    """Export article as Word document."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        filename = f"{query}_{timestamp}.docx"
        filepath = os.path.join(output_dir, filename)
        
        doc = Document()
        
        # Parse markdown-style article (simple parsing)
        lines = article.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Title (# heading)
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Heading 2 (## heading)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            # Heading 3 (### heading)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            # Regular paragraph
            else:
                doc.add_paragraph(line)
        
        # Add sources section
        if sources:
            doc.add_page_break()
            doc.add_heading('Sources', level=2)
            for i, url in enumerate(sources, 1):
                doc.add_paragraph(f"{i}. {url}", style='List Number')
        
        doc.save(filepath)
        return filepath
        
    except ImportError:
        print("Warning: python-docx not installed, falling back to markdown")
        return _export_markdown(article, sources, query, timestamp, output_dir)
    except Exception as e:
        print(f"Warning: DOCX export error: {e}, falling back to markdown")
        return _export_markdown(article, sources, query, timestamp, output_dir)


def _export_notion(article: str, sources: list, query: str) -> str:
    """Export article to Notion (placeholder for future implementation)."""
    print("Warning: Notion export not yet implemented, falling back to markdown")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return _export_markdown(article, sources, query, timestamp, "outputs")
